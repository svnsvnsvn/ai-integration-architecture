"""
Document Processing Pipeline with LangGraph
A multi-step workflow for secure document ingestion, processing, and storage
"""

import os
import json
import sqlite3
import hashlib
import asyncio
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, TypedDict
from dataclasses import dataclass, asdict

# Document processing
import PyPDF2
from docx import Document as DocxDocument

# LangGraph and LangChain
from langgraph.graph import StateGraph, END
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# AI Providers
from langchain_openai import ChatOpenAI
from langchain_groq import ChatGroq
from langchain_ollama import OllamaLLM

# Environment and validation
from dotenv import load_dotenv
from pydantic import BaseModel, Field

# Load environment
load_dotenv()

# State definition for LangGraph
class DocumentState(TypedDict):
    # Input
    file_path: str
    file_type: str
    
    # Processing state
    raw_content: str
    chunks: List[str]
    metadata: Dict[str, Any]
    processed_content: Dict[str, Any]
    
    # Output
    document_id: str
    status: str
    error: Optional[str]
    
    # Security
    content_hash: str
    safe_content: bool

@dataclass
class ProcessedDocument:
    """Processed document structure"""
    document_id: str
    file_path: str
    file_type: str
    content_hash: str
    summary: str
    key_points: List[str]
    metadata: Dict[str, Any]
    chunk_count: int
    processed_at: datetime
    safe_content: bool

def initialize_llm():
    """Initialize LLM based on environment configuration"""
    ai_provider = os.getenv("AI_PROVIDER", "ollama").lower()
    
    try:
        if ai_provider == "openai":
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                print("⚠️  OpenAI API key not found. Falling back to Ollama.")
                return OllamaLLM(model="llama3.1")
            return ChatOpenAI(
                model="gpt-3.5-turbo",
                api_key=api_key,
                temperature=0.3
            )
        
        elif ai_provider == "groq":
            api_key = os.getenv("GROQ_API_KEY")
            if not api_key:
                print("⚠️  Groq API key not found. Falling back to Ollama.")
                return OllamaLLM(model="llama3.1")
            return ChatGroq(
                model="llama3-8b-8192",
                api_key=api_key,
                temperature=0.3
            )
        
        else:  # ollama (default)
            model = os.getenv("OLLAMA_MODEL", "llama3.1")
            print(f"🦙 Using Ollama with model: {model}")
            print("💡 Make sure Ollama is running: 'ollama serve'")
            return OllamaLLM(model=model)
            
    except Exception as e:
        print(f"⚠️  Error initializing {ai_provider}: {e}")
        print("🦙 Falling back to Ollama")
        return OllamaLLM(model="llama3.1")

class DocumentProcessor:
    """Core document processing logic"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        # Initialize real LLM
        self.llm = initialize_llm()
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract PDF text: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract DOCX text: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            raise Exception(f"Failed to extract TXT text: {str(e)}")
    
    def is_content_safe(self, content: str) -> bool:
        """Basic content safety check"""
        # Simple checks for potentially unsafe content
        unsafe_patterns = [
            "password", "secret", "api_key", "private_key",
            "credit_card", "ssn", "social_security"
        ]
        
        content_lower = content.lower()
        for pattern in unsafe_patterns:
            if pattern in content_lower:
                return False
        return True
    
    def generate_content_hash(self, content: str) -> str:
        """Generate SHA-256 hash of content"""
        return hashlib.sha256(content.encode()).hexdigest()
    
    def generate_summary(self, text: str) -> str:
        """Generate a summary using real AI"""
        try:
            prompt = f"""Please provide a concise summary of the following document in 1-2 sentences:

{text[:2000]}  

Summary:"""
            
            response = self.llm.invoke(prompt)
            if hasattr(response, 'content'):
                return response.content.strip()
            return str(response).strip()
            
        except Exception as e:
            print(f"⚠️  Summary generation failed: {e}")
            return f"Document summary unavailable. Content length: {len(text)} characters."
    
    def extract_key_points(self, text: str) -> List[str]:
        """Extract key points using real AI"""
        try:
            prompt = f"""Extract 3-5 key points from this document. Format as a simple list:

{text[:2000]}

Key points:
1."""
            
            response = self.llm.invoke(prompt)
            content = response.content if hasattr(response, 'content') else str(response)
            
            # Parse the response into a list
            points = []
            for line in content.split('\n'):
                line = line.strip()
                if line and (line.startswith(('1.', '2.', '3.', '4.', '5.', '-', '•'))):
                    # Clean up the point
                    point = line.split('.', 1)[-1].strip() if '.' in line else line[1:].strip()
                    if point:
                        points.append(point)
            
            return points[:5] if points else ["Key points extraction unavailable"]
            
        except Exception as e:
            print(f"⚠️  Key points extraction failed: {e}")
            return ["Key points extraction unavailable due to processing error"]
    
    def analyze_content_safety(self, text: str) -> bool:
        """Analyze content safety using real AI"""
        try:
            prompt = f"""Analyze this document for potentially sensitive information. 
Look for: passwords, API keys, personal data, confidential information, private details.

Document excerpt:
{text[:1000]}

Respond with just "SAFE" or "UNSAFE" followed by a brief reason."""

            response = self.llm.invoke(prompt)
            content = response.content if hasattr(response, 'content') else str(response)
            
            # Simple parsing - look for SAFE/UNSAFE
            content_lower = content.lower()
            if 'unsafe' in content_lower:
                return False
            elif 'safe' in content_lower:
                return True
            else:
                # Fallback to keyword detection
                return self.is_content_safe(text)
                
        except Exception as e:
            print(f"⚠️  AI safety analysis failed: {e}")
            # Fallback to basic keyword detection
            return self.is_content_safe(text)

class DocumentDatabase:
    """SQLite database for storing processed documents"""
    
    def __init__(self, db_path: str = "documents.db"):
        self.db_path = db_path
        self.init_database()
    
    def init_database(self):
        """Initialize the database schema"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                document_id TEXT PRIMARY KEY,
                file_path TEXT NOT NULL,
                file_type TEXT NOT NULL,
                content_hash TEXT NOT NULL,
                summary TEXT,
                key_points TEXT,
                metadata TEXT,
                chunk_count INTEGER,
                processed_at TEXT,
                safe_content BOOLEAN
            )
        """)
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS document_chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                document_id TEXT,
                chunk_index INTEGER,
                content TEXT,
                FOREIGN KEY (document_id) REFERENCES documents (document_id)
            )
        """)
        
        conn.commit()
        conn.close()
    
    def save_document(self, doc: ProcessedDocument, chunks: List[str]):
        """Save processed document to database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            # Save document metadata
            cursor.execute("""
                INSERT OR REPLACE INTO documents 
                (document_id, file_path, file_type, content_hash, summary, 
                 key_points, metadata, chunk_count, processed_at, safe_content)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                doc.document_id,
                doc.file_path,
                doc.file_type,
                doc.content_hash,
                doc.summary,
                json.dumps(doc.key_points),
                json.dumps(doc.metadata),
                doc.chunk_count,
                doc.processed_at.isoformat(),
                doc.safe_content
            ))
            
            # Delete existing chunks
            cursor.execute("DELETE FROM document_chunks WHERE document_id = ?", (doc.document_id,))
            
            # Save chunks
            for i, chunk in enumerate(chunks):
                cursor.execute("""
                    INSERT INTO document_chunks (document_id, chunk_index, content)
                    VALUES (?, ?, ?)
                """, (doc.document_id, i, chunk))
            
            conn.commit()
            
        except Exception as e:
            conn.rollback()
            raise e
        finally:
            conn.close()
    
    def get_document(self, document_id: str) -> Optional[ProcessedDocument]:
        """Retrieve a document by ID"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute("SELECT * FROM documents WHERE document_id = ?", (document_id,))
        row = cursor.fetchone()
        
        if row:
            return ProcessedDocument(
                document_id=row[0],
                file_path=row[1],
                file_type=row[2],
                content_hash=row[3],
                summary=row[4],
                key_points=json.loads(row[5]) if row[5] else [],
                metadata=json.loads(row[6]) if row[6] else {},
                chunk_count=row[7],
                processed_at=datetime.fromisoformat(row[8]),
                safe_content=bool(row[9])
            )
        
        conn.close()
        return None
    
    def list_documents(self) -> List[Dict[str, Any]]:
        """List all documents"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT document_id, file_path, file_type, summary, processed_at, safe_content 
            FROM documents ORDER BY processed_at DESC
        """)
        
        docs = []
        for row in cursor.fetchall():
            docs.append({
                "document_id": row[0],
                "file_path": row[1],
                "file_type": row[2],
                "summary": row[3],
                "processed_at": row[4],
                "safe_content": bool(row[5])
            })
        
        conn.close()
        return docs

# LangGraph Node Functions
def ingest_document(state: DocumentState) -> DocumentState:
    """Step 1: Ingest and extract text from document"""
    try:
        processor = DocumentProcessor()
        file_path = state["file_path"]
        file_type = state["file_type"].lower()
        
        print(f"Ingesting document: {file_path}")
        
        # Extract text based on file type
        if file_type == "pdf":
            raw_content = processor.extract_text_from_pdf(file_path)
        elif file_type == "docx":
            raw_content = processor.extract_text_from_docx(file_path)
        elif file_type == "txt":
            raw_content = processor.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Generate content hash and check safety
        content_hash = processor.generate_content_hash(raw_content)
        print("🔍 Analyzing content safety with AI...")
        safe_content = processor.analyze_content_safety(raw_content)
        
        # Update state
        state.update({
            "raw_content": raw_content,
            "content_hash": content_hash,
            "safe_content": safe_content,
            "status": "ingested",
            "document_id": f"doc_{content_hash[:8]}"
        })
        
        print(f"Document ingested: {len(raw_content)} characters")
        return state
        
    except Exception as e:
        state.update({
            "status": "error",
            "error": str(e)
        })
        print(f"Ingestion failed: {e}")
        return state

def process_content(state: DocumentState) -> DocumentState:
    """Step 2: Process and analyze document content"""
    try:
        if state["status"] == "error":
            return state
        
        print(f"Processing content for document: {state['document_id']}")
        
        processor = DocumentProcessor()
        
        # Split into chunks
        documents = [Document(page_content=state["raw_content"])]
        chunks = processor.text_splitter.split_documents(documents)
        chunk_texts = [chunk.page_content for chunk in chunks]
        
        # Generate real AI summary
        print("🤖 Generating AI summary...")
        summary = processor.generate_summary(state['raw_content'])
        
        # Extract real AI key points
        print("🔍 Extracting key points...")
        key_points = processor.extract_key_points(state['raw_content'])
        
        # Create metadata
        metadata = {
            "file_size": len(state["raw_content"]),
            "word_count": len(state["raw_content"].split()),
            "chunk_count": len(chunk_texts),
            "has_sensitive_content": not state["safe_content"]
        }
        
        # Update state
        state.update({
            "chunks": chunk_texts,
            "processed_content": {
                "summary": summary,
                "key_points": key_points
            },
            "metadata": metadata,
            "status": "processed"
        })
        
        print(f"Content processed: {len(chunk_texts)} chunks created")
        return state
        
    except Exception as e:
        state.update({
            "status": "error",
            "error": str(e)
        })
        print(f"Processing failed: {e}")
        return state

def store_document(state: DocumentState) -> DocumentState:
    """Step 3: Store processed document in database"""
    try:
        if state["status"] == "error":
            return state
        
        print(f"Storing document: {state['document_id']}")
        
        # Create processed document object
        processed_doc = ProcessedDocument(
            document_id=state["document_id"],
            file_path=state["file_path"],
            file_type=state["file_type"],
            content_hash=state["content_hash"],
            summary=state["processed_content"]["summary"],
            key_points=state["processed_content"]["key_points"],
            metadata=state["metadata"],
            chunk_count=len(state["chunks"]),
            processed_at=datetime.now(),
            safe_content=state["safe_content"]
        )
        
        # Save to database
        db = DocumentDatabase()
        db.save_document(processed_doc, state["chunks"])
        
        state.update({
            "status": "completed"
        })
        
        print(f"Document stored successfully")
        return state
        
    except Exception as e:
        state.update({
            "status": "error",
            "error": str(e)
        })
        print(f"Storage failed: {e}")
        return state

def handle_error(state: DocumentState) -> DocumentState:
    """Handle errors in the pipeline"""
    print(f"Error handling for document: {state.get('document_id', 'unknown')}")
    print(f"   Error: {state.get('error', 'unknown error')}")
    
    # Log error (in production, use proper logging)
    error_log = {
        "timestamp": datetime.now().isoformat(),
        "document_id": state.get("document_id"),
        "file_path": state.get("file_path"),
        "error": state.get("error"),
        "status": state.get("status")
    }
    
    with open("error_log.json", "a") as f:
        f.write(json.dumps(error_log) + "\n")
    
    return state

# Create the LangGraph workflow
def create_document_processing_workflow():
    """Create the document processing workflow graph"""
    
    workflow = StateGraph(DocumentState)
    
    # Add nodes
    workflow.add_node("ingest", ingest_document)
    workflow.add_node("process", process_content)
    workflow.add_node("store", store_document)
    workflow.add_node("error_handler", handle_error)
    
    # Set entry point
    workflow.set_entry_point("ingest")
    
    # Add conditional edges
    def should_continue_after_ingest(state: DocumentState) -> str:
        if state["status"] == "error":
            return "error_handler"
        return "process"
    
    def should_continue_after_process(state: DocumentState) -> str:
        if state["status"] == "error":
            return "error_handler"
        return "store"
    
    def should_continue_after_store(state: DocumentState) -> str:
        if state["status"] == "error":
            return "error_handler"
        return END
    
    # Add edges
    workflow.add_conditional_edges("ingest", should_continue_after_ingest)
    workflow.add_conditional_edges("process", should_continue_after_process)
    workflow.add_conditional_edges("store", should_continue_after_store)
    workflow.add_edge("error_handler", END)
    
    return workflow.compile()

# Main pipeline class
class DocumentPipeline:
    """Main document processing pipeline"""
    
    def __init__(self):
        self.workflow = create_document_processing_workflow()
        self.db = DocumentDatabase()
    
    async def process_document(self, file_path: str) -> DocumentState:
        """Process a single document through the pipeline"""
        
        # Validate file exists
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type
        file_type = path.suffix.lower().lstrip('.')
        if file_type not in ['pdf', 'docx', 'txt']:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Initial state
        initial_state = DocumentState(
            file_path=str(path.absolute()),
            file_type=file_type,
            raw_content="",
            chunks=[],
            metadata={},
            processed_content={},
            document_id="",
            status="pending",
            error=None,
            content_hash="",
            safe_content=True
        )
        
        print(f"Starting document processing pipeline for: {file_path}")
        print("="*60)
        
        # Run the workflow
        result = self.workflow.invoke(initial_state)
        
        print("="*60)
        print(f"Pipeline completed with status: {result['status']}")
        
        return result
    
    def list_processed_documents(self) -> List[Dict[str, Any]]:
        """List all processed documents"""
        return self.db.list_documents()
    
    def get_document_details(self, document_id: str) -> Optional[ProcessedDocument]:
        """Get detailed information about a processed document"""
        return self.db.get_document(document_id)

if __name__ == "__main__":
    # This will be used by the demo script
    pass
